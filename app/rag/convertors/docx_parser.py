from io import BytesIO
import uuid

from docx import Document

from app.rag.parsers.base import (
    BaseParser,
    StructuredDocument,
    Asset,
)


class DocxParser(BaseParser):

    def parse(
            self,
            file_bytes: bytes,
    ) -> StructuredDocument:

        doc = Document(
            BytesIO(file_bytes)
        )

        markdown_parts = []

        text_parts = []

        assets = []

        # =========================
        # 段落解析
        # =========================

        for para in doc.paragraphs:

            text = para.text.strip()

            if not text:
                continue

            style_name = para.style.name.lower()

            # heading识别
            if "heading 1" in style_name:
                md = f"# {text}"

            elif "heading 2" in style_name:
                md = f"## {text}"

            elif "heading 3" in style_name:
                md = f"### {text}"

            else:
                md = text

            markdown_parts.append(md)

            text_parts.append(text)

        # =========================
        # 图片提取
        # =========================

        rels = doc.part.rels

        for rel in rels.values():

            if "image" not in rel.target_ref:
                continue

            image_data = rel.target_part.blob

            content_type = rel.target_part.content_type

            ext = content_type.split("/")[-1]

            filename = (
                f"{uuid.uuid4()}.{ext}"
            )

            asset = Asset(
                asset_id=str(uuid.uuid4()),
                filename=filename,
                mime_type=content_type,
                data=image_data,
            )

            assets.append(asset)

            markdown_parts.append(
                f"![{filename}](assets/{filename})"
            )

        markdown = "\n\n".join(markdown_parts)

        text = "\n".join(text_parts)

        return StructuredDocument(
            text=text,
            markdown=markdown,
            assets=assets,
            metadata={
                "source_type": "docx"
            },
        )
